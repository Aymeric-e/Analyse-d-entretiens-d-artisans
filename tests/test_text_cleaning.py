"""
Unit tests for text cleaning module.

Tests the InterviewCleaner class from src/preprocessing/text_cleaning.py.
Includes tests for:
- clean_artisan_text: Validates text cleaning with parametrized test cases
- process_file: Tests .docx file processing with minimal setup

Fixtures:
    tmp_path: Pytest fixture providing a temporary directory for test files.
"""

import pytest
from docx import Document

from preprocessing.text_cleaning import InterviewCleaner


@pytest.mark.parametrize(
    "input_text, expected",
    [
        ("[Interviewé] – Bonjour, je m'appelle Paul.", "Bonjour, je m'appelle Paul."),
        ("(01:23) Salut tout le monde", "Salut tout le monde"),
        ("  (04:48) ", ""),
        ("[Interviewe]   -  Voilà   un   test", "Voilà un test"),
    ],
)
def test_clean_artisan_text_examples(input_text, expected):
    """
    Test clean_artisan_text with various input patterns.

    Validates that speaker tags, timestamps, and excessive whitespace are removed
    while preserving core text content.

    Args:
        input_text: Raw text with interview speaker markers and timestamps.
        expected: Expected cleaned output.
    """
    cleaner = InterviewCleaner()
    cleaned = cleaner.clean_artisan_text(input_text)
    assert cleaned == expected


def test_process_file_minimum_words(tmp_path):
    """
    Test process_file method with a minimal .docx document.

    Creates a temporary .docx file containing one interviewee paragraph and one
    researcher paragraph, then validates that only the interviewee content is
    extracted and returned as a list of cleaned text records.

    Args:
        tmp_path: Pytest fixture providing a temporary directory for test files.
    """

    doc = Document()
    doc.add_paragraph("[Interviewé] Bonjour tout le monde.")
    doc.add_paragraph("[Chercheur] Question")
    p = tmp_path / "test.docx"
    doc.save(p)

    cleaner = InterviewCleaner()
    rows = cleaner.process_file(p)
    # Should extract the interviewee line and ignore the researcher line
    assert isinstance(rows, list)
    assert len(rows) == 1
    assert rows[0]["text"].startswith("Bonjour")
